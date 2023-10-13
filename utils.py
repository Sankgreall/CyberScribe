from docx import Document
import tiktoken
from PyPDF2 import PdfReader
import mimetypes
import os

def read_text_file(file_path):
    with open(file_path, 'r') as file:
        return file.read()
    
def is_audio_file(file_path):
    mime_type, _ = mimetypes.guess_type(file_path)
    return mime_type.startswith('audio') if mime_type else False

def read_docx_file(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def write_to_file(filename, content, folder="output"):
    """Writes content to a file in the ./output directory."""
    
    # Create the output directory if it doesn't exist
    if not os.path.exists(f"./{folder}"):
        os.makedirs(f"./{folder}")
        
    # Create the complete path to the file
    filepath = os.path.join(f"./{folder}", filename)
    
    # Write content to the file
    with open(filepath, 'w') as f:
        f.write(content)

def count_tokens(text):
    """Returns the number of tokens in a text string."""
    encoding_name = "cl100k_base"
    encoding = tiktoken.get_encoding(encoding_name)
    num_tokens = len(encoding.encode(text))
    return num_tokens

def extract_text_from_pdf(pdf_path):
    reader = PdfReader(pdf_path)
    text = ""
    for page in reader.pages:
        text += page.extract_text()
    return text

def extract_text_from_docx(docx_path):
    doc = Document(docx_path)
    return "\n".join([para.text for para in doc.paragraphs])

def retry_error_callback(retry_state):
    print("Received too many errors, quitting...")
    
    # Directly get the exception from the outcome
    exc = retry_state.outcome.exception()
    if exc:
        print("Underlying Exception:", type(exc))
        print("Exception Details:", exc)
    
    raise RuntimeError("Too many retries due to errors.")